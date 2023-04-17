import logging
import re
from collections import namedtuple
from pathlib import Path
from typing import List, Optional

import gender_guesser.detector as gender
import pycountry
from deep_translator import GoogleTranslator
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from gender_guesser.detector import NoCountryError

from invoicer.config import Config
from invoicer.order import Address, Customer, Invoice, Item, Order
from invoicer.order_mail_parsers import find_country_index
from invoicer.utils import eur, get_short_date, get_year, prepend_zeros


Font = namedtuple("Font", ("name", "size"))
TableIndices = namedtuple("TableIndices", ("items", "sum", "passport"))
Replacement = namedtuple("Replacement", ("old", "new", "font"))


class InvoiceGenerator:
    def __init__(self, cfg: Config, template_path: Path) -> None:
        self.default_font = Font("Arial", 11)
        self.table_indices = TableIndices(items=0, sum=1, passport=2)
        self.config = cfg
        self.template_path = template_path

    def _create_invoice_nr(self, year_invoice: int):
        try:
            with open(".last_invoice_count", "r", encoding="utf-8") as f:
                last_invoice_count = int(f.read())
        except FileNotFoundError:
            last_invoice_count = self.config.invoiceCountStart

        invoice_count = str(last_invoice_count + 1)
        invoice_count = prepend_zeros(value=invoice_count, num_min_digits=3) 

        invoice_nr = str(year_invoice) + invoice_count

        with open(".last_invoice_count", "w", encoding="utf-8") as f:
            f.write(str(invoice_count))

        return invoice_nr


    def _check_errors(self, order: Order, init_default=True, dump_errors: Optional[bool] = None, dump_errors_path: Optional[Path] = None):
        order = Order(**order.__dict__)

        to_be_checked_attrs = (
            ("date", ""),
            ("invoice", Invoice(address=Address(full_name="", address=""))),
            ("items", []),
            ("shipping_cost", 0),
            ("payment_method", ""),
            ("customer", Customer(name="", surname=""))
        )
        errors = []
        def _check(name, default):
            if getattr(order, name) is None:
                logging.error(f"{name} is None")
                errors.append(name)
                if init_default:
                    setattr(order, name, default)

        for to_be_checked_attr in to_be_checked_attrs:
            _check(name=to_be_checked_attr[0], default=to_be_checked_attr[1])
        
        if dump_errors and len(errors) > 0:
            with open(file=dump_errors_path, mode="w+", encoding="utf-8") as f:
                f.write("\n".join(errors))

        return (order, errors)

    def generate(self, order: Order):
        """
        Validate order and log errors. 
        Replace missing data accordingly.
        Generate invoice from order.
        """
        if order.invoice.date is None:
            order.invoice.date = get_short_date(order.date)
        if order.invoice.number is None:
            order.invoice.number = self._create_invoice_nr(
                year_invoice=get_year(order.invoice.date)
            )

        order, errors = self._check_errors(
            order=order,
            init_default=True,
            dump_errors=True,
            dump_errors_path=Path(f"Invoice-{order.invoice.number}-Errors.txt")
        )
        
        invoice = Document(self.template_path)

        self._replace_paragraphs(order=order, invoice=invoice)
        self._replace_tables(order=order, invoice=invoice)

        docx_path = Path(f"docs/Invoice-{order.invoice.number}.docx")
        invoice.save(docx_path)
        # pdf_path = docx_path.with_suffix(".pdf")
        # subprocess.run(
        #     [
        #         "libreoffice",
        #         "--convert-to",
        #         "pdf",
        #         str(docx_path),
        #         "--outdir",
        #         str(pdf_path.parent),
        #     ]
        # )
        return (docx_path, errors)

    def _replace_paragraphs(self, order: Order, invoice=Document):
        full_address = _get_full_address(order)

        replacements = (
            Replacement(r"{{address}}", full_address, Font("Courier New", 10)),
            Replacement(r"{{date}}", order.invoice.date, Font("Arial", 11)),
            Replacement(r"{{invoice_nr}}", order.invoice.number, Font("Arial", 16)),
            Replacement(r"{{salutation}}", _guess_salutation(order), Font("Arial", 32)),
            Replacement(r"{{payment_method}}", order.payment_method, Font("Arial", 8)),
        )

        for replacement in replacements:
            self._replace(**replacement._asdict(), invoice=invoice)

    def _replace_tables(self, order: Order, invoice: Document):
        shipping = Item(
            count=1,
            description="Verpackung & Lieferung",
            price=order.shipping_cost,
            unit="psch.",
            tax_rate=0.19,
        )
        items = order.items[:] + [shipping]

        item_table = invoice.tables[self.table_indices.items]
        _replace_item_table(t=item_table, items=items)

        sum_table = invoice.tables[self.table_indices.sum]
        _replace_sum_table(t=sum_table, items=items)

        passport_table = invoice.tables[self.table_indices.passport]
        _replace_passport_table(t=passport_table, order=order)

        tables = (item_table, sum_table, passport_table)

        for table in tables:
            _change_font(table, Font("Calibri", 10))

    def _replace(self, invoice: Document, old: str, new: str, font: Optional[Font] = None):
        # if font is None:
        #     font = self.default_font
        index = self._get_index(token=old, invoice=invoice)
        p = invoice.paragraphs[index]
        paragraph_replace_text(paragraph=p, regex=re.compile(old), replace_str=new)
        # # TODO: This changes font for the whole paragraph, convert to run.

        # if font:
        #     _change_paragraph_font(p=p, new=font)
        # replaced = invoice.paragraphs[index].text.replace(old, new)
        # invoice.paragraphs[index].text = replaced

    def _get_index(self, token: str, invoice: Document):
        paragraphs = [p.text for p in invoice.paragraphs]
        for i, p in enumerate(paragraphs):
            if token in p:
                return i


def _guess_salutation(order: Order):
    address = _get_full_address(order)
    address_parts = address.split("\n")
    full_name = address_parts[0]

    country = address_parts[-1]
    d = gender.Detector()
    name_parts = full_name.split(" ")
    name_pick = name_parts[0]
    if name_parts[0] in ["Dr.", "Dr", "Prof.", "Prof"]:
        name_pick = name_parts[1]

    try:
        result = d.get_gender(name=name_pick, country=country.lower())
    except NoCountryError:
        result = d.get_gender(name=name_pick)

    if order.customer.surname in full_name:
        saluatation_name = order.customer.surname
    else:
        saluatation_name = full_name

    if result in ["male", "mostly_male"]:
        return f"Sehr geehrter Herr {saluatation_name}"

    if result in ["female", "mostly_female"]:
        return f"Sehr geehrte Frau {saluatation_name}"

    return f"Sehr geehrte(r) Frau/Herr {saluatation_name}"


def _write_item_to_row(index, item, row):
    row.cells[0].text = str(index)
    row.cells[1].text = str(item.count)
    row.cells[2].text = item.unit
    row.cells[3].text = item.description
    row.cells[4].text = eur(item.unit_price_net)
    row.cells[5].text = eur(item.unit_price_gross)
    row.cells[6].text = eur(item.total_price_gross)


def _replace_item_table(t, items: List[Item]):
    # t.style.font.name = "Calibri"
    # t.style.font.size = Pt(10)
    for i, item in enumerate(items):
        row = t.add_row()
        _write_item_to_row(index=i + 1, item=item, row=row)

    align_column(t.columns[0], WD_ALIGN_PARAGRAPH.CENTER)
    align_column(t.columns[1], WD_ALIGN_PARAGRAPH.CENTER)
    align_column(t.columns[2], WD_ALIGN_PARAGRAPH.CENTER)
    align_column(t.columns[3], WD_ALIGN_PARAGRAPH.LEFT)
    align_column(t.columns[4], WD_ALIGN_PARAGRAPH.RIGHT)
    align_column(t.columns[5], WD_ALIGN_PARAGRAPH.RIGHT)
    align_column(t.columns[6], WD_ALIGN_PARAGRAPH.RIGHT)


def _replace_sum_table(t, items: List[Item]):
    def get_sum(tax_rate: float, attr: str):
        return sum([getattr(item, attr) for item in items if item.tax_rate == tax_rate])

    # t.style.font.name = "Calibri"
    # t.style.font.size = Pt(10)

    sum_net_price_7 = get_sum(tax_rate=0.07, attr="total_price_net")
    sum_net_price_19 = get_sum(tax_rate=0.19, attr="total_price_net")
    tax_7 = get_sum(tax_rate=0.07, attr="tax")
    tax_19 = get_sum(tax_rate=0.19, attr="tax")

    values = (sum_net_price_7, sum_net_price_19, tax_7, tax_19)
    for i, value in enumerate(values):
        t.cell(i, 6).text = eur(value)
    t.cell(len(values), 6).text = eur(sum(values))

    align_column(t.columns[6], WD_ALIGN_PARAGRAPH.RIGHT)

    if len(items) - 1 == 1:
        s_e_7 = "1"
    else:
        s_e_7 = f"{1}-{len(items) - 1}"

    t.cell(2, 3).text = t.cell(2, 3).text.replace(r"{{7_s_e}}", s_e_7)
    t.cell(3, 3).text = t.cell(3, 3).text.replace(r"{{19_s_e}}", f"{len(items)}")
    align_column(t.columns[3], WD_ALIGN_PARAGRAPH.RIGHT)


def _replace_passport_table(t, order: Order):
    # t.style.font.name = "Calibri"
    # t.style.font.size = Pt(10)
    items = order.items
    t.cell(7, 4).text = "\n".join(map(lambda item: item.description, items)) + "\n"
    date = order.invoice.date
    invoice_nr = order.invoice.number
    t.cell(10, 4).text = "".join(reversed(date.split("."))) + "-" + invoice_nr[4:]


def _get_full_address(order: Order):
    address = order.invoice.address
    full_name = address.full_name
    address = address.address

    address_parts = address.split("\n")
    country_index = find_country_index(address_parts)
    country = address_parts[country_index]

    if country == "Deutschland":
        address_parts = address_parts[:2]
    else:
        country_en = GoogleTranslator(source="de", target="en").translate(country)
        address_parts[country_index] = country_en

    address = "\n".join(address_parts)
    return "\n".join([full_name, address])


def _get_country_en(country_de: str):
    for country in pycountry.countries:
        if _(country.name) == country_de:
            return country.name

    # if country == "Deutschland":
    #     residential_address_end_index = 2
    # else:
    #     invoice_address[country_index] = GoogleTranslator(source="de", target="en").translate(country)


def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph


def _change_font(table, font):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.name = font.name
                    run.font.size = Pt(font.size)


def align_column(column, alignment):
    for cell in column.cells:
        for p in cell.paragraphs:
            p.alignment = alignment


def _get_separate_fields(fields: str):
    fields = [f"{{{{{field}}}}}" for field in fields.split(", ")]
    return fields